from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from pathlib import Path


ROOT = Path(__file__).resolve().parents[1]
OUT_DOCX = ROOT / "NAUB_CSPMS_Chapters_1_to_5.docx"
OUT_APPENDIX = ROOT / "APPENDIX_CODE_SUMMARIES.md"


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_doc_defaults(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.first_line_indent = Inches(0.5)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(12)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.first_line_indent = Inches(0)

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1)


def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run = p.add_run()
    run._r.append(field_begin)
    run._r.append(instr)
    run._r.append(field_end)


def add_center_heading(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Inches(0)
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    return p


def add_section_heading(doc, number, title):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.first_line_indent = Inches(0)
    r1 = p.add_run(f"{number}\t")
    r1.bold = False
    r2 = p.add_run(title)
    r2.bold = True
    return p


def add_subheading(doc, text):
    p = doc.add_paragraph(style="Heading 3")
    p.paragraph_format.first_line_indent = Inches(0)
    p.add_run(text).bold = True
    return p


def add_para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style=None)
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.left_indent = Inches(0.25)
        p.add_run(f"{item}")


def add_table(doc, number, title, headers, rows, note=None):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0)
    r = p.add_run(f"Table {number}")
    r.bold = True
    t = doc.add_paragraph()
    t.paragraph_format.first_line_indent = Inches(0)
    rt = t.add_run(title)
    rt.italic = True
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True)
        shade_cell(hdr[i], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
    if note:
        n = doc.add_paragraph()
        n.paragraph_format.first_line_indent = Inches(0)
        rn = n.add_run("Note. ")
        rn.italic = True
        n.add_run(note)


def add_figure_placeholder(doc, number, title, placeholder, note=None):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0)
    r = p.add_run(f"Figure {number}")
    r.bold = True
    t = doc.add_paragraph()
    t.paragraph_format.first_line_indent = Inches(0)
    rt = t.add_run(title)
    rt.italic = True
    ph = doc.add_paragraph()
    ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ph.paragraph_format.first_line_indent = Inches(0)
    rr = ph.add_run(f"({placeholder})")
    rr.italic = True
    if note:
        n = doc.add_paragraph()
        n.paragraph_format.first_line_indent = Inches(0)
        rn = n.add_run("Note. ")
        rn.italic = True
        n.add_run(note)


def add_reference(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p.paragraph_format.left_indent = Inches(0.5)
    p.add_run(text)


doc = Document()
set_doc_defaults(doc)
add_page_number(doc.sections[0])

add_center_heading(doc, "INTELLIGENT AI-DRIVEN STUDENTS' PROJECT MANAGEMENT SYSTEM FOR NIGERIAN ARMY UNIVERSITY BIU")
for line in [
    "A Final Year Project Report",
    "Department of Computer Science",
    "Nigerian Army University Biu",
    "Prepared in APA 7th Edition Style",
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Inches(0)
    p.add_run(line)

doc.add_page_break()
add_center_heading(doc, "TABLE OF CONTENTS")
toc = [
    ("1.0", "CHAPTER ONE: INTRODUCTION"),
    ("1.1", "Background of the Study"),
    ("1.2", "Statement of the Problem"),
    ("1.3", "Aim and Objectives of the Study"),
    ("1.4", "The Scope of the Study"),
    ("1.5", "Significance of the Study"),
    ("1.6", "Definition of Terms"),
    ("2.0", "CHAPTER TWO: LITERATURE REVIEW"),
    ("2.1", "Introduction"),
    ("2.2", "Review of Existing Model"),
    ("2.3", "Theoretical Framework"),
    ("2.4", "Conceptual Review"),
    ("2.5", "Review of Related Works"),
    ("2.6", "Summary of Literature Review"),
    ("3.0", "CHAPTER THREE: SYSTEM ANALYSIS AND DESIGN"),
    ("3.1", "Specify the Methodology Adopted"),
    ("3.2", "Analysis of the Existing System"),
    ("3.3", "The Proposed System Design"),
    ("3.4", "The Architecture, Model, etc."),
    ("3.5", "Design Flowchart/Algorithm"),
    ("3.6", "The Advantage of the New System"),
    ("4.0", "CHAPTER FOUR: SYSTEM IMPLEMENTATION"),
    ("4.1", "Results"),
    ("4.2", "Discussion"),
    ("5.0", "CHAPTER FIVE: SUMMARY, CONCLUSION AND RECOMMENDATION"),
    ("5.1", "Summary"),
    ("5.2", "Conclusion"),
    ("5.3", "Limitation of the Study"),
    ("5.4", "Recommendation"),
    ("5.5", "Contribution to Knowledge"),
]
for n, title in toc:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0)
    p.add_run(n).bold = n.endswith(".0")
    p.add_run("\t")
    p.add_run(title).bold = n.endswith(".0")

doc.add_page_break()

# Chapter One
add_center_heading(doc, "1.0\tCHAPTER ONE: INTRODUCTION")
add_section_heading(doc, "1.1", "Background of the Study")
for text in [
    "Final-year projects are a major academic requirement through which undergraduate students demonstrate independent inquiry, disciplinary competence, problem-solving ability, and research communication skills. In computer science and related disciplines, the final-year project also allows students to translate theoretical learning into a practical system, algorithm, or research solution that can be evaluated in relation to a real problem.",
    "At Nigerian Army University Biu (NAUB), project supervision involves students, supervisors, project coordinators, and departmental management. The process requires topic selection, supervision assignment, chapter submission, supervisor feedback, revisions, deadline management, and final reporting. When these activities are handled manually, important records can be scattered across paper files, informal chat messages, and individual devices. This makes monitoring difficult and slows down decision-making.",
    "The growth of web-based information systems has made it possible to centralize academic workflows. A computerized students' project management system can store project records, support online submission, preserve feedback history, and give coordinators visibility into project progress. For a university that emphasizes technology, research, and development, such a system is not merely convenient; it supports the institutional goal of improving academic administration through digital tools.",
    "The system implemented in this study is called NAUB Prism. It is an intelligent AI-driven students' project management system built as a web application. The implementation uses Next.js, TypeScript, PostgreSQL, Prisma, Auth.js, UploadThing, Pusher, and a rule-based AI Project Guide. These technologies replace the earlier conceptual assumption of a PHP/MySQL/LAMP implementation and more accurately represent the system that was developed.",
    "The intelligent component is the Project Guide chatbot. It does not replace lecturers or supervisors. Instead, it provides immediate academic guidance on routine research questions such as chapter structure, topic narrowing, objectives, methodology, literature review organization, academic writing, and APA 7 referencing. This allows supervisors to focus more on substantive academic review while students receive faster support for common project-related questions.",
]:
    add_para(doc, text)

add_section_heading(doc, "1.2", "Statement of the Problem")
for text in [
    "The existing undergraduate project supervision process in many departments depends on physical submissions, face-to-face meetings, and informal messaging platforms. This arrangement creates delays because supervisors may not always be available to review printed documents or respond to routine questions. It also makes it difficult to maintain a complete record of submissions, feedback, deadlines, and student progress.",
    "A second problem is the absence of structured visibility for coordinators. Without a central dashboard, coordinators may not know which students are unassigned, which projects are active, which submissions are awaiting review, and which milestones are overdue. This limits the ability of the department to intervene early when students are falling behind.",
    "A third problem is that ordinary web-based project portals often digitize document upload but do not provide intelligent student support. Students still depend fully on supervisors for questions about proposal writing, chapter structure, research methodology, and referencing. This dependence delays routine guidance and increases the workload of supervisors.",
    "Therefore, this study addresses the need for a role-based, web-based, intelligent students' project management system that centralizes project supervision, supports online submission and feedback, provides progress monitoring, enables secure communication, and offers a rule-based AI Project Guide for common academic guidance.",
]:
    add_para(doc, text)

add_section_heading(doc, "1.3", "Aim and Objectives of the Study")
add_subheading(doc, "Aim")
add_para(doc, "The aim of this study is to design and implement an intelligent AI-driven students' project management system for Nigerian Army University Biu that improves undergraduate project supervision through centralized project tracking, online document submission, supervisor feedback, coordinator monitoring, secure communication, and AI-assisted academic guidance.")
add_subheading(doc, "Specific Objectives")
for item in [
    "1. Examine the current undergraduate project management process and identify the challenges associated with manual and informal supervision practices.",
    "2. Determine the functional and non-functional requirements for an intelligent students' project management system.",
    "3. Design the system architecture, database schema, user interfaces, and workflow models for NAUB Prism.",
    "4. Implement the system using Next.js, TypeScript, PostgreSQL, Prisma, Auth.js, UploadThing, Pusher, and a rule-based AI Project Guide.",
    "5. Evaluate the implemented modules in terms of functionality, usability, security, and alignment with the needs of students, supervisors, and coordinators.",
    "6. Document the implementation, testing outcomes, limitations, and possible future enhancements of the system.",
]:
    add_para(doc, item)

add_section_heading(doc, "1.4", "The Scope of the Study")
for text in [
    "This study is limited to the design and implementation of NAUB Prism for undergraduate final-year project supervision at Nigerian Army University Biu. The implemented system supports three operational roles: student, supervisor, and coordinator. Students can register, create project records, view milestones, submit progress with documents, receive feedback, chat with assigned supervisors, view notifications, and use the Project Guide. Supervisors can manage assigned students, review submissions, provide feedback, set deadlines, and communicate with students. Coordinators can manage users, assign supervisors, create milestone templates, monitor analytics, publish announcements, and view audit logs.",
    "The study covers document upload for PDF and DOCX project files, structured feedback, milestone tracking, in-app notifications, realtime chat, dashboard analytics, and a rule-based academic guidance chatbot. It does not cover postgraduate research management, full plagiarism detection, automatic grading, external institutional system integration, payment systems, or live deployment across all NAUB departments.",
]:
    add_para(doc, text)

add_section_heading(doc, "1.5", "Significance of the Study")
for text in [
    "The study is significant to students because it provides a single platform for tracking project progress, submitting drafts, receiving feedback, viewing deadlines, and obtaining instant academic guidance. This reduces uncertainty and helps students manage their final-year project work more systematically.",
    "The study is significant to supervisors because it gives them a structured review environment. Instead of depending on scattered documents and informal messages, supervisors can view submissions, record feedback, monitor assigned students, and communicate through the platform.",
    "The study is significant to coordinators because it provides administrative visibility. Coordinators can manage students and supervisors, assign supervisees, monitor project activity, view system analytics, and maintain an audit trail of important actions.",
    "The study is also significant to the university because it demonstrates how a modern web-based system can support academic administration, reduce paper-based inefficiencies, and promote technology-driven project supervision.",
]:
    add_para(doc, text)

add_section_heading(doc, "1.6", "Definition of Terms")
terms = [
    ("Artificial Intelligence", "The use of computational logic to provide responses or assistance that resembles intelligent support. In NAUB Prism, AI refers to the rule-based Project Guide that gives academic guidance from embedded research knowledge."),
    ("AI Project Guide", "A rule-based academic assistant within NAUB Prism that responds to student questions on chapter structure, research methodology, academic writing, and APA 7 conventions."),
    ("Auth.js", "The authentication framework used in the implemented system to manage credentials login and role-aware sessions."),
    ("Computerized Students' Project Management System", "A web-based platform that supports project registration, submission, supervision, feedback, communication, monitoring, and reporting."),
    ("Coordinator", "A departmental user responsible for managing users, assignments, milestone templates, announcements, analytics, and audit oversight."),
    ("Milestone", "A defined project stage or deliverable, such as proposal, Chapter One, methodology, or final submission, used to monitor progress."),
    ("Next.js", "The React-based framework used to implement the web application, server-rendered pages, API routes, and server actions."),
    ("PostgreSQL", "The relational database used to store users, projects, submissions, feedback, conversations, notifications, deadlines, and audit logs."),
    ("Prisma", "The object-relational mapper used to define the schema and interact with the PostgreSQL database."),
    ("Role-Based Access Control", "A security mechanism that restricts system functions according to the user's assigned role: student, supervisor, or coordinator."),
]
for term, definition in terms:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(f"{term}: ").bold = True
    p.add_run(definition)

doc.add_page_break()

# Chapter Two
add_center_heading(doc, "2.0\tCHAPTER TWO: LITERATURE REVIEW")
add_section_heading(doc, "2.1", "Introduction")
add_para(doc, "This chapter reviews literature related to computerized students' project management systems, web-based academic administration, role-based access control, database-driven applications, realtime communication, and AI-assisted academic guidance. The review is organized to support the design and implementation of NAUB Prism as an intelligent project supervision platform.")

add_section_heading(doc, "2.2", "Review of Existing Model")
for text in [
    "Existing research supervision systems generally focus on digitizing the submission and feedback process. Typical features include student registration, supervisor assignment, document upload, review comments, and administrator reporting. These systems improve record keeping when compared with paper-based methods, but many still depend entirely on supervisors for routine academic guidance.",
    "A related web-based research supervision model implemented with Laravel and MySQL provided student, supervisor, and administrator roles. The model supported communication and document management, but its limitation was the absence of intelligent support. Students still waited for supervisors to answer common questions about research structure and formatting.",
    "NAUB Prism extends this model by combining project management functions with a Project Guide chatbot. The implemented system also differs technically because it uses Next.js, PostgreSQL, Prisma, and server actions rather than PHP/MySQL. This architectural choice improves modularity, type safety, and maintainability while retaining the academic supervision goals of earlier systems.",
]:
    add_para(doc, text)

add_section_heading(doc, "2.3", "Theoretical Framework")
add_subheading(doc, "Systems Development Life Cycle")
add_para(doc, "The Systems Development Life Cycle provides a structured approach to software development through stages such as planning, requirements analysis, design, implementation, testing, and maintenance. This study follows the SDLC logic because the system was developed from identified supervision problems, translated into requirements, designed as a role-based architecture, implemented as a web application, and evaluated through module-level functionality.")
add_subheading(doc, "Technology Acceptance Model")
add_para(doc, "The Technology Acceptance Model explains users' acceptance of information systems through perceived usefulness and perceived ease of use. NAUB Prism is intended to be useful because it reduces manual supervision problems, and easy to use because each role receives a dashboard with only the functions relevant to that user.")
add_subheading(doc, "DeLone and McLean Information Systems Success Model")
add_para(doc, "The DeLone and McLean model provides a basis for evaluating information system success using system quality, information quality, service quality, use, user satisfaction, and net benefits. These dimensions apply to NAUB Prism because the system must be secure, reliable, accurate, accessible, and beneficial to students, supervisors, and coordinators.")

add_section_heading(doc, "2.4", "Conceptual Review")
for title, text in [
    ("Project Management in Undergraduate Research", "Undergraduate project management involves planning, supervision, deliverables, deadlines, feedback, revision, and final reporting. In this study, project management is implemented through projects, milestones, submissions, deadlines, and status tracking."),
    ("Web-Based Academic Systems", "A web-based academic system allows users to access institutional services through browsers without installing desktop software. NAUB Prism uses this model so that students, supervisors, and coordinators can access the platform from standard desktop or mobile browsers."),
    ("Role-Based Access Control", "Role-based access control protects academic data by allowing each user to access only the functions relevant to their role. In NAUB Prism, students cannot access coordinator analytics, supervisors cannot manage all users, and coordinators cannot read private student-supervisor chat content unless they are direct participants."),
    ("AI-Assisted Academic Guidance", "AI-assisted academic guidance refers to computer-supported responses that help students understand research tasks. In this implementation, the Project Guide is rule-based and knowledge-driven. It supports students with general academic guidance but does not generate complete projects or replace supervisor judgment."),
]:
    add_subheading(doc, title)
    add_para(doc, text)

add_section_heading(doc, "2.5", "Review of Related Works")
add_table(
    doc,
    "2.1",
    "Summary of Related Systems and Gaps",
    ["Author/Source", "System Focus", "Strength", "Gap Addressed by NAUB Prism"],
    [
        ["Abiola et al. (2024)", "Student project management system", "Digitized project records", "Limited intelligent guidance"],
        ["Kolade et al. (2024)", "Final-year project supervision platform", "Reduced paper-based submission", "Routine academic support still supervisor-dependent"],
        ["Musty Joda (2024)", "Interactive research supervision", "Improved communication and document handling", "Recommended AI enhancement"],
        ["NAUB Prism", "AI-driven project supervision", "Combines project tracking, feedback, chat, notifications, and Project Guide", "Implemented for NAUB undergraduate supervision"],
    ],
    "The table summarizes the practical gap that informed the implemented system: digitization alone is useful, but intelligent guidance and role-based progress monitoring strengthen the supervision process.",
)

add_section_heading(doc, "2.6", "Summary of Literature Review")
add_para(doc, "The reviewed literature shows that computerized project management systems improve the supervision process by reducing paper dependence, preserving records, and supporting electronic submission. However, many existing systems do not include intelligent academic support. NAUB Prism responds to this gap by implementing a role-based web platform with project tracking, document submission, feedback, realtime communication, coordinator analytics, and a rule-based Project Guide.")

doc.add_page_break()

# Chapter Three
add_center_heading(doc, "3.0\tCHAPTER THREE: SYSTEM ANALYSIS AND DESIGN")
add_section_heading(doc, "3.1", "Specify the Methodology Adopted")
for text in [
    "This study adopted the Structured Systems Analysis and Design Methodology (SSADM) within a broader Systems Development Life Cycle approach. SSADM was suitable because the system requirements were clear: project registration, user authentication, submission, feedback, notification, communication, progress monitoring, and AI-assisted guidance.",
    "The methodology involved analysis of the existing system, identification of functional and non-functional requirements, database design, interface design, system architecture design, implementation, and testing. The implemented technologies were selected based on maintainability, security, and suitability for modern web development.",
]:
    add_para(doc, text)

add_section_heading(doc, "3.2", "Analysis of the Existing System")
for text in [
    "The existing system is largely manual. Students prepare project topics and chapter drafts, supervisors review submitted work, and coordinators monitor progress through informal updates. Communication often occurs through physical meetings and messaging applications. This creates delays, weak accountability, and poor institutional visibility.",
    "The main problems are delayed feedback, document loss, lack of centralized progress tracking, limited topic visibility, absence of structured milestone monitoring, and no instant support for common research questions. These weaknesses justify the development of a computerized system.",
]:
    add_para(doc, text)

add_section_heading(doc, "3.3", "The Proposed System Design")
add_para(doc, "The proposed system, NAUB Prism, is a role-based web application for managing undergraduate research supervision. It provides dashboards for students, supervisors, and coordinators. The system allows students to create project records, upload project documents, submit progress, view feedback, monitor milestones, receive notifications, and use the Project Guide. Supervisors review submissions, give feedback, set deadlines, and communicate with students. Coordinators manage users, supervisor assignments, milestone templates, announcements, analytics, and audit logs.")
add_table(
    doc,
    "3.1",
    "Functional Requirements of NAUB Prism",
    ["Requirement", "Description", "Implemented Module"],
    [
        ["Authentication", "Secure login and registration with role-based routing", "Auth.js, middleware, session helpers"],
        ["Project Management", "Student project profile and milestone tracking", "Projects actions and project pages"],
        ["Submission Management", "Upload and submit project files for review", "Submission actions, UploadThing"],
        ["Feedback", "Supervisor decision and structured comments", "Feedback form and server action"],
        ["Communication", "Realtime role-aware chat", "Chat actions and Pusher"],
        ["Notifications", "In-app alerts for submissions, feedback, assignments, and deadlines", "Notification actions and Pusher"],
        ["Coordinator Oversight", "User management, assignments, analytics, audit logs", "Coordinator dashboard and actions"],
        ["AI Project Guide", "Rule-based academic support for research questions", "Project guide API and knowledge base"],
    ],
)
add_table(
    doc,
    "3.2",
    "Non-Functional Requirements",
    ["Requirement", "Design Response"],
    [
        ["Usability", "Responsive role dashboards with focused menus and reusable UI components."],
        ["Security", "Credential authentication, hashed passwords, middleware route protection, and server-side role checks."],
        ["Reliability", "PostgreSQL relational storage, Prisma schema constraints, and audit logging."],
        ["Performance", "Server-rendered pages, cached session access, and fallback handling for some dashboard reads."],
        ["Maintainability", "Feature-based code organization under src/features and reusable infrastructure under src/lib."],
        ["Scalability", "Modular role model, database-backed records, and deployment-ready Next.js architecture."],
    ],
)

add_section_heading(doc, "3.4", "The Architecture, Model, etc.")
add_para(doc, "NAUB Prism uses a three-tier architecture. The presentation tier contains the browser-based user interface built with Next.js, React, Tailwind CSS, and reusable UI components. The application tier contains server-rendered pages, server actions, API routes, authentication logic, validation, notification logic, and the Project Guide API. The data tier contains PostgreSQL accessed through Prisma.")
add_figure_placeholder(doc, "3.1", "Three-Tier Architecture of NAUB Prism", "Insert architecture diagram showing browser/user interface, Next.js application/server actions/API routes, and PostgreSQL/Prisma data layer.", "The architecture reflects the implemented system and replaces the earlier PHP/MySQL/LAMP assumption.")
add_figure_placeholder(doc, "3.2", "Use Case Diagram of NAUB Prism", "Insert use case diagram showing Student, Supervisor, Coordinator, and Project Guide interactions.", "The implemented roles are student, supervisor, and coordinator.")
add_figure_placeholder(doc, "3.3", "Entity Relationship Diagram of NAUB Prism", "Insert ERD showing User, ResearchProject, ProjectMilestone, Submission, SubmissionFile, Feedback, Conversation, Message, Notification, Announcement, Deadline, and AuditLog.", "The ERD should be generated from the Prisma schema.")
add_table(
    doc,
    "3.3",
    "Core Database Entities",
    ["Entity", "Purpose"],
    [
        ["User", "Stores account details, roles, status, student-supervisor assignment, and profile information."],
        ["ResearchProject", "Stores each student's project title, topic, abstract, session year, and supervisor link."],
        ["ProjectMilestone", "Tracks project stages, due dates, and milestone status."],
        ["Submission", "Stores submitted work, version, review status, and related milestone."],
        ["Feedback", "Stores supervisor comments and review decisions."],
        ["Conversation and Message", "Store role-aware academic communication."],
        ["Notification", "Stores in-app alerts for key events."],
        ["AuditLog", "Stores high-level activity metadata for accountability."],
    ],
)

add_section_heading(doc, "3.5", "Design Flowchart/Algorithm")
add_para(doc, "The main system logic follows role-based authentication and action-specific authorization. A user signs in, the system verifies credentials, middleware checks the user's role, and the appropriate dashboard is displayed. Each operation is then validated and authorized on the server before database changes are made.")
add_figure_placeholder(doc, "3.4", "General System Flowchart", "Insert flowchart showing login, role verification, dashboard routing, module access, validation, database operation, notification, and audit logging.")
add_subheading(doc, "Submission Review Algorithm")
for item in [
    "1. Student logs into the system.",
    "2. System confirms that the user role is STUDENT.",
    "3. Student selects a milestone or general progress update.",
    "4. Student uploads PDF/DOCX files and submits the form.",
    "5. Server validates the form and file payload.",
    "6. System checks that the student has a project and an assigned supervisor.",
    "7. System creates the submission record and related file records.",
    "8. System updates milestone status where applicable.",
    "9. System notifies the supervisor and records an audit log.",
    "10. Supervisor reviews the submission and submits feedback.",
]:
    add_para(doc, item)

add_section_heading(doc, "3.6", "The Advantage of the New System")
for text in [
    "The new system centralizes project supervision records and reduces dependence on printed submissions. It improves accountability by preserving submission history, feedback, notifications, and audit logs.",
    "It improves communication through role-aware chat and realtime notifications. Students can communicate with assigned supervisors, while coordinators can communicate with students or supervisors without being allowed to inspect unrelated private student-supervisor conversations.",
    "It improves academic support by integrating a Project Guide that gives instant responses to routine project questions. This reduces repetitive supervisor workload while preserving the supervisor's role in academic judgment and formal feedback.",
    "It improves administrative visibility by giving coordinators analytics on students, supervisors, active projects, submissions, pending reviews, unassigned students, and recent activity.",
]:
    add_para(doc, text)

doc.add_page_break()

# Chapter Four
add_center_heading(doc, "4.0\tCHAPTER FOUR: SYSTEM IMPLEMENTATION")
add_section_heading(doc, "4.1", "RESULTS")
add_para(doc, "This chapter presents the implementation results of NAUB Prism. The system was implemented as a browser-based application using Next.js 15, React 19, TypeScript, Tailwind CSS, PostgreSQL, Prisma, Auth.js, UploadThing, Pusher, and Inngest. The implemented modules are presented with their interface placeholders and summary code references. Full source code extracts are summarized in the appendix.")

add_table(
    doc,
    "4.1",
    "Implementation Environment",
    ["Component", "Technology Used", "Purpose"],
    [
        ["Frontend and Routing", "Next.js App Router, React, TypeScript", "Pages, layouts, UI rendering, and route organization"],
        ["Styling", "Tailwind CSS, custom UI components", "Responsive academic dashboard interface"],
        ["Database", "PostgreSQL with Prisma ORM", "Relational data persistence"],
        ["Authentication", "Auth.js credentials with JWT sessions", "Login, session, and role-based access"],
        ["File Upload", "UploadThing", "PDF/DOCX project document upload"],
        ["Realtime", "Pusher", "Chat and live notifications"],
        ["Background Jobs", "Inngest", "Deadline reminders and weekly pending review digest"],
        ["AI Guide", "Rule-based knowledge API", "Academic guidance chatbot"],
    ],
)

modules = [
    ("4.1", "Login and Authentication Interface", "Insert Screenshot: login page showing email/password fields and role-aware sign-in result.", "The login module authenticates users through Auth.js credentials. Passwords are checked using bcrypt, and authenticated users are routed to the appropriate dashboard according to role."),
    ("4.2", "Student Dashboard Interface", "Insert Screenshot: student dashboard showing project summary, milestones, deadlines, announcements, and recent submissions.", "The student dashboard summarizes the student's project status, submission counts, milestone completion, supervisor assignment, deadlines, and announcements."),
    ("4.3", "Project Profile and Milestone Interface", "Insert Screenshot: student project page showing project details and milestone progress.", "The project module allows students to create or update project details. When a project is first created, active milestone templates can be used to generate project milestones."),
    ("4.4", "Submission Upload Interface", "Insert Screenshot: new submission form showing title, milestone selector, progress summary, and document upload area.", "The submission module accepts progress details and PDF/DOCX files. Server-side logic verifies the student's role, project existence, supervisor assignment, and validated file metadata before creating a submission."),
    ("4.5", "Supervisor Review and Feedback Interface", "Insert Screenshot: supervisor review page showing submitted document metadata and feedback form.", "The supervisor review module allows supervisors to approve, reject, or request revision. Feedback decisions update both submission status and related milestone status."),
    ("4.6", "Coordinator Dashboard Interface", "Insert Screenshot: coordinator dashboard showing analytics cards, submission outcome chart, unassigned students, and recent audit logs.", "The coordinator dashboard provides institutional visibility into project supervision activity without exposing private document contents or unrelated chat transcripts."),
    ("4.7", "Supervisor Assignment Interface", "Insert Screenshot: coordinator assignment page showing student-supervisor assignment controls.", "The assignment module links students to supervisors and updates both the user record and related project record. Notifications are sent to both parties after assignment."),
    ("4.8", "Realtime Chat Interface", "Insert Screenshot: chat panel showing academic conversation between assigned users.", "The chat module creates private conversations between permitted users. Students can message assigned supervisors, supervisors can message assigned students, and coordinators can message students or supervisors."),
    ("4.9", "Notification Interface", "Insert Screenshot: notifications page showing unread and read notifications.", "The notification module creates in-app alerts for submissions, feedback, messages, deadlines, assignments, announcements, and system events."),
    ("4.10", "AI Project Guide Interface", "Insert Screenshot: floating Project Guide chat window showing academic question and generated response.", "The Project Guide module uses embedded knowledge and rule matching to provide academic guidance on research structure, methodology, writing, and APA 7 conventions."),
]
for number, title, placeholder, explanation in modules:
    add_figure_placeholder(doc, number, title, placeholder)
    add_para(doc, explanation)

add_table(
    doc,
    "4.2",
    "Summary of Implemented Modules and Source Locations",
    ["Module", "Main Source File(s)", "Implementation Summary"],
    [
        ["Authentication", "src/lib/auth.ts; src/features/auth/actions.ts; src/middleware.ts", "Credentials login, JWT session, role routing, inactive user restriction."],
        ["Project Management", "src/features/projects/actions.ts", "Project creation/update, milestone templates, custom milestones, deadlines."],
        ["Submission and Feedback", "src/features/submissions/actions.ts", "Submission creation, file metadata storage, supervisor feedback, status updates."],
        ["Coordinator Management", "src/features/users/actions.ts", "Profile update, user creation, status update, supervisor assignment."],
        ["Chat", "src/features/chat/actions.ts; src/components/shared/chat-panel.tsx", "Conversation creation, message authorization, realtime updates."],
        ["Notifications", "src/lib/notifications.ts; src/features/notifications/actions.ts", "In-app notification creation, read/unread handling, realtime trigger."],
        ["AI Project Guide", "src/app/api/ai/project-guide/route.ts; src/lib/ai/project-guide-knowledge.ts", "Rule-based academic response generation."],
        ["Database", "prisma/schema.prisma; src/lib/prisma.ts", "PostgreSQL schema, Prisma models, relational integrity."],
    ],
)

add_table(
    doc,
    "4.3",
    "Functional Test Results",
    ["Test Case", "Expected Result", "Observed Result", "Status"],
    [
        ["Student login with valid credentials", "Student dashboard opens", "Role dashboard is loaded", "Passed"],
        ["Supervisor attempts student dashboard route", "Redirect to supervisor home", "Middleware prevents access", "Passed"],
        ["Student submits project without supervisor", "Submission blocked", "System returns assignment error", "Passed"],
        ["Student submits valid document", "Submission record created", "Submission and file metadata saved", "Passed"],
        ["Supervisor reviews assigned submission", "Feedback saved and student notified", "Status and notification updated", "Passed"],
        ["Coordinator assigns supervisor", "Student and supervisor records updated", "Assignment notification created", "Passed"],
        ["User sends chat message", "Message appears in conversation", "Pusher realtime event triggered", "Passed"],
        ["Student asks Project Guide question", "Academic guidance returned", "Rule-based answer displayed", "Passed"],
    ],
)

add_section_heading(doc, "4.2", "DISCUSSION")
for text in [
    "The implemented system achieved the major functional requirements identified in Chapter Three. Authentication and authorization were implemented at multiple levels through middleware, session helpers, and server actions. This design reduces the possibility of unauthorized access because each critical operation checks the user's role before performing database changes.",
    "The student workflow was successfully implemented through project creation, milestone tracking, document upload, submission creation, notification, and feedback viewing. This directly addresses the manual submission problem because documents and feedback can now be linked to structured database records.",
    "The supervisor workflow was implemented through assigned student views, review queues, feedback forms, deadlines, and chat. This improves review accountability because each decision is stored with the supervisor identity, date, and submission relationship.",
    "The coordinator workflow was implemented through user management, supervisor assignments, analytics, milestone templates, announcements, search, and audit logs. This gives the coordinator a system-wide administrative view while preserving privacy boundaries around chat and research documents.",
    "The Project Guide was implemented as a rule-based academic assistant rather than a generative AI model connected to an external API. This is a controlled implementation choice. It allows the system to provide reliable guidance on chapter structure, APA 7, academic writing, and methodology without inventing references or replacing supervisor judgment. However, it also means that the guide is limited to the knowledge and rules embedded in the application.",
    "Overall, the implementation shows that a modern web-based architecture can effectively support the undergraduate project supervision process. The system improves organization, traceability, communication, and student support while maintaining a modular codebase that can be extended in future work.",
]:
    add_para(doc, text)

doc.add_page_break()

# Chapter Five
add_center_heading(doc, "5.0\tCHAPTER FIVE: SUMMARY, CONCLUSION AND RECOMMENDATION")
add_section_heading(doc, "5.1", "Summary")
for text in [
    "This study designed and implemented NAUB Prism, an intelligent AI-driven students' project management system for undergraduate final-year project supervision at Nigerian Army University Biu. The system was motivated by the limitations of manual project supervision, including delayed feedback, scattered records, poor progress visibility, document loss, and overdependence on supervisors for routine academic questions.",
    "Chapter One introduced the study, stated the problem, aim, objectives, scope, significance, and operational definitions. Chapter Two reviewed relevant literature on computerized project management systems, web-based academic administration, role-based access control, and AI-assisted academic guidance. Chapter Three analyzed the existing system and presented the proposed architecture, requirements, database model, use cases, and workflow design. Chapter Four described the implementation results and discussed the major modules.",
    "The implemented system includes authentication, role-based dashboards, project records, milestones, submissions, file upload, supervisor feedback, coordinator management, realtime chat, notifications, deadlines, audit logs, analytics, and a rule-based AI Project Guide.",
]:
    add_para(doc, text)

add_section_heading(doc, "5.2", "Conclusion")
add_para(doc, "The study concludes that a computerized and intelligent project management system can improve undergraduate project supervision by centralizing records, structuring submission and feedback workflows, improving coordinator visibility, supporting realtime academic communication, and providing instant guidance for routine research questions. NAUB Prism demonstrates that modern web technologies such as Next.js, PostgreSQL, Prisma, Auth.js, UploadThing, and Pusher can be combined to produce a maintainable, secure, and functional supervision platform. The AI Project Guide adds value by assisting students with academic guidance, although it remains a support tool and not a replacement for lecturers or supervisors.")

add_section_heading(doc, "5.3", "Limitation of the Study")
for text in [
    "The system was developed and evaluated within the scope of undergraduate final-year project supervision and was not integrated with official university portals, payment systems, plagiarism detection systems, or institutional email infrastructure.",
    "The Project Guide is rule-based and depends on embedded knowledge. It does not use a live large language model, does not learn from student interactions, and does not store chatbot interaction history for later analysis.",
    "The implemented database stores department as a user attribute rather than as a full department entity with independent departmental rules, programmes, or approval workflows.",
    "The system does not yet include a dedicated defense scheduling module, topic similarity detection, or full administrative role separation beyond the coordinator role.",
    "Testing was limited to functional and module-level validation. Broader usability testing with a large number of real students, supervisors, and coordinators was not conducted within the study period.",
]:
    add_para(doc, text)

add_section_heading(doc, "5.4", "Recommendation")
for item in [
    "1. The system should be extended with a dedicated defense scheduling module that allows coordinators to create panels, venues, dates, and defense reports.",
    "2. A topic similarity detection module should be added to flag potentially duplicated or highly similar project topics before approval.",
    "3. The Project Guide should be improved by storing chatbot interactions, tracking unanswered questions, and expanding the knowledge base with department-approved project guidelines.",
    "4. A full administrator role should be introduced for institution-level configuration while coordinators remain responsible for departmental supervision workflows.",
    "5. The system should be integrated with institutional email or SMS services for stronger notification delivery if approved by the university.",
    "6. Further usability testing should be conducted with real students, supervisors, and coordinators to measure satisfaction, ease of use, and supervision efficiency.",
]:
    add_para(doc, item)

add_section_heading(doc, "5.5", "Contribution to Knowledge")
for text in [
    "This study contributes to knowledge by demonstrating a practical implementation of an intelligent project supervision platform for a Nigerian university context. It shows how project management, academic supervision, realtime communication, role-based access control, and AI-assisted academic guidance can be combined in a single system.",
    "The study also contributes a modern technical model for CSPMS development using Next.js, PostgreSQL, Prisma, server actions, realtime notifications, and a rule-based academic assistant. This differs from many earlier PHP/MySQL project portals and provides a maintainable architecture for future student project systems.",
    "Finally, the study clarifies the boundary between AI support and academic supervision. The Project Guide assists with routine academic guidance but does not replace the professional judgment, feedback, and approval responsibilities of supervisors.",
]:
    add_para(doc, text)

doc.add_page_break()
add_center_heading(doc, "REFERENCES")
refs = [
    "American Psychological Association. (2020). Publication manual of the American Psychological Association (7th ed.). American Psychological Association.",
    "Davis, F. D. (1989). Perceived usefulness, perceived ease of use, and user acceptance of information technology. MIS Quarterly, 13(3), 319-340.",
    "DeLone, W. H., & McLean, E. R. (2003). The DeLone and McLean model of information systems success: A ten-year update. Journal of Management Information Systems, 19(4), 9-30.",
    "Laudon, K. C., & Laudon, J. P. (2020). Management information systems: Managing the digital firm (16th ed.). Pearson.",
    "Pressman, R. S., & Maxim, B. R. (2020). Software engineering: A practitioner's approach (9th ed.). McGraw-Hill.",
    "Project Management Institute. (2021). A guide to the project management body of knowledge (PMBOK guide) (7th ed.). Project Management Institute.",
    "Trent University Academic Skills. (n.d.). APA 7 style formatting guidelines. https://www.trentu.ca/academicskills/documentation-guide/apa-style/apa-7-style-formatting-guidelines",
    "University of St. Thomas Libraries. (2025). APA citation style 7th ed.: Paper formatting. https://libguides.stthomas.edu/APAstyle/formatting",
    "Metro State University. (n.d.). APA 7th edition quick guide. https://www.metrostate.edu/academics/success/tutoring/resources/apa-quick-guide",
]
for ref in refs:
    add_reference(doc, ref)

doc.save(OUT_DOCX)

appendix = """# Appendix: Summary of Key Source Code Extracts

This appendix summarizes key implementation files used in NAUB Prism. The full files remain in the project repository, while these summaries identify the code sections that may be referenced in the final project appendix.

## Appendix A: Authentication and Role-Based Session Control

**Source files:** `src/lib/auth.ts`, `src/lib/session.ts`, `src/middleware.ts`, `src/features/auth/actions.ts`

**Summary:** These files implement credentials-based login, JWT session handling, password verification, and role-based route protection. `auth.ts` configures Auth.js with a credentials provider. `session.ts` exposes `requireUser()` to protect server pages and actions. `middleware.ts` redirects unauthenticated users and prevents students, supervisors, and coordinators from accessing dashboards outside their roles.

**Academic relevance:** This code supports the security requirement of role-based access control and ensures that each user category can only perform permitted actions.

## Appendix B: Project, Milestone, and Deadline Management

**Source files:** `src/features/projects/actions.ts`, `prisma/schema.prisma`

**Summary:** The project module allows students to create or update their research project records. It also creates milestones from active milestone templates and allows supervisors/coordinators to add custom milestones and deadlines. The Prisma schema defines `ResearchProject`, `MilestoneTemplate`, `ProjectMilestone`, and `Deadline` relationships.

**Academic relevance:** This code implements the project tracking requirement and replaces manual tracking of project stages with structured database-backed milestones.

## Appendix C: Submission and Supervisor Feedback Workflow

**Source files:** `src/features/submissions/actions.ts`, `src/components/shared/submission-form.tsx`, `src/components/shared/feedback-form.tsx`, `src/lib/uploadthing.ts`

**Summary:** Students submit progress updates with optional milestone links and uploaded PDF/DOCX files. The server action validates the submission, checks that the student has a project and assigned supervisor, stores file metadata, updates milestone status, notifies the supervisor, and writes an audit log. Supervisors review submissions and save feedback decisions such as approved, rejected, or needs revision.

**Academic relevance:** This code supports the central objective of moving chapter submission and supervisor feedback from paper/informal chat into a traceable system.

## Appendix D: Realtime Chat and Notification System

**Source files:** `src/features/chat/actions.ts`, `src/components/shared/chat-panel.tsx`, `src/lib/notifications.ts`, `src/features/notifications/actions.ts`

**Summary:** The chat module creates authorized conversations between students, supervisors, and coordinators. Pusher is used to deliver realtime messages and notification events. The notification module stores in-app alerts for submissions, feedback, messages, assignments, and deadlines. Email sending is optional and skipped when the email provider key is not configured.

**Academic relevance:** This code addresses communication delays by providing secure, role-aware academic messaging and event notifications.

## Appendix E: AI Project Guide

**Source files:** `src/app/api/ai/project-guide/route.ts`, `src/lib/ai/project-guide-knowledge.ts`, `src/components/shared/project-guide-chat.tsx`

**Summary:** The Project Guide is a rule-based academic assistant. The client component displays a floating chat interface, sends user questions to the API route, and renders markdown responses. The knowledge file contains academic guidance on chapter structure, methodology, literature review, APA 7, academic writing, and supervisor feedback. The route normalizes recent chat messages and returns matched guidance.

**Academic relevance:** This code implements the intelligent support feature that differentiates NAUB Prism from ordinary web-based project submission systems.
"""

OUT_APPENDIX.write_text(appendix, encoding="utf-8")

print(f"Wrote {OUT_DOCX}")
print(f"Wrote {OUT_APPENDIX}")
